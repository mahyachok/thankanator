import csv 
import docx

with open('donor_list.csv', newline= '') as f:
    reader = csv.reader(f)

    doc = docx.Document("Donation Receipt.docx")

    for row in reader:
        donors = list(row)
        #name the charecteristic of the words from the list
        last_name= donors[0]
        first_name= donors[1]
        address= donors[2]
        donation= donors[3]
        #split the address into two parts
        
        
        #if there isnt a comma or periood in something then it will default it to the except
        try:
           
            if  '.' in address:
                address =address.split('.', 1)

            elif "," in address:
                address =address.split(",", 1)

            addressp1=address[0]
            addressp2=address[1]

        except:#if it dosent pare the way i like this is sort of the default thing. not really ideal. 
            addressp1=donors[2]
            addressp2=""
        
        #this is where every thing is in the section of the word document, run is used to specify the text in a paragraph
        doc.paragraphs[5].text= first_name, last_name #first name
        doc.paragraphs[6].text= addressp1#address
        doc.paragraphs[7].text=addressp2 #address part2
        doc.paragraphs[10].runs[0].text= "Dear", first_name,"," #first
        doc.paragraphs[12].runs[1].text= "$",donation #donations

        #this saves the document with the name of the person
        doc.save(f"{last_name}{first_name}.docx")
    